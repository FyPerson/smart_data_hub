#!/usr/bin/env python3
"""
md2docx.py - Markdown → Word 转换脚本
用法: python scripts/md2docx.py input.md [output.docx]
      如不指定 output，自动生成同名 .docx

流程: pandoc 转换 → python-docx 后处理（表格美化、引用块调整、分隔线删除、排版规范化）
依赖: pandoc 3.x, python-docx
模板: reference.docx（由 build_reference.py 生成）

排版规范（参考 GB/T 9704 + 企业汇报实践）：
  页面: A4, 上下 2.54cm, 左右 3.17cm
  字体: 全局微软雅黑
  正文: 小四(12pt), 行距 28磅, 首行缩进 2字符
  标题: H1=16pt H2=15pt H3=14pt H4=12pt, 全部加粗黑色
  表格: 蓝色表头白字 + 斑马纹 + 浅灰边框
  引用: 破折号开头 + 斜体灰色 + 左缩进
"""

import sys
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 配置 ---
PANDOC = r"C:\Users\FY\AppData\Local\Microsoft\WinGet\Packages\JohnMacFarlane.Pandoc_Microsoft.Winget.Source_8wekyb3d8bbwe\pandoc-3.9\pandoc.exe"
REFERENCE_DOC = str(Path(__file__).parent / "reference.docx")
BLUE = "2C5AA0"
FONT_CN = "微软雅黑"


def step1_pandoc_convert(md_path: str, docx_path: str):
    """pandoc 基础转换

    pandoc 按当前工作目录解析 md 里的相对路径（图片引用等），
    这里强制切换到 md 所在目录，确保无论从哪里运行脚本都能正确找到图片。
    """
    md_abs = os.path.abspath(md_path)
    docx_abs = os.path.abspath(docx_path)
    md_dir = os.path.dirname(md_abs)
    md_name = os.path.basename(md_abs)
    cmd = [PANDOC, md_name, "-o", docx_abs, "--reference-doc", REFERENCE_DOC]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=md_dir)
    if result.returncode != 0:
        print(f"pandoc 错误: {result.stderr}")
        sys.exit(1)
    print(f"  [1/5] pandoc 转换完成")


def step2_beautify_tables(doc: Document):
    """表格美化：蓝色表头白字、斑马纹、紧凑行距、自适应列宽"""
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)

        # 宽度 100%
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn("w:type"), "pct")
            tblW.set(qn("w:w"), "5000")

        # 列宽：按内容长度智能分配
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is not None:
            tblPr.remove(tblLayout)
        tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))
        _auto_column_widths(table)

        # 表格边框：浅灰细线
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        BORDER = 'w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"'
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top {BORDER}/>'
            f'<w:left {BORDER}/>'
            f'<w:bottom {BORDER}/>'
            f'<w:right {BORDER}/>'
            f'<w:insideH {BORDER}/>'
            f'<w:insideV {BORDER}/>'
            f'</w:tblBorders>'
        ))

        for row_idx, row in enumerate(table.rows):
            # 表头行：设置跨页重复显示
            if row_idx == 0:
                trPr = row._tr.find(qn("w:trPr"))
                if trPr is None:
                    trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
                    row._tr.insert(0, trPr)
                tblHeader = trPr.find(qn("w:tblHeader"))
                if tblHeader is None:
                    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

            num_cols = len(row.cells)
            for col_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                    cell._tc.insert(0, tcPr)

                # 紧凑内边距
                tcMar = tcPr.find(qn("w:tcMar"))
                if tcMar is not None:
                    tcPr.remove(tcMar)
                tcPr.append(parse_xml(
                    f'<w:tcMar {nsdecls("w")}>'
                    f'<w:top w:w="40" w:type="dxa"/>'
                    f'<w:left w:w="80" w:type="dxa"/>'
                    f'<w:bottom w:w="40" w:type="dxa"/>'
                    f'<w:right w:w="80" w:type="dxa"/>'
                    f"</w:tcMar>"
                ))

                if row_idx == 0:
                    # 表头：蓝底白字加粗 + 居中 + 不换行
                    _set_cell_shading(tcPr, BLUE)
                    _format_cell_runs(cell, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
                    # 不换行
                    noWrap = tcPr.find(qn("w:noWrap"))
                    if noWrap is None:
                        tcPr.append(parse_xml(f'<w:noWrap {nsdecls("w")}/>'))
                    # 居中对齐
                    for para in cell.paragraphs:
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # 偶数行：浅灰背景
                    if row_idx % 2 == 0:
                        _set_cell_shading(tcPr, "F5F7FA")
                    _format_cell_runs(cell, color=RGBColor(0x33, 0x33, 0x33), bold=False)

                # 紧凑行距
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.line_spacing = 1.2

    print(f"  [2/5] 表格美化完成（{len(doc.tables)} 个表格）")


def step3_fix_blockquotes(doc: Document):
    """引用块：去掉网页装饰，改为缩进+斜体+灰色左竖线"""
    count = 0
    for para in doc.paragraphs:
        if para.style.name == "Block Text":
            count += 1
            pPr = para._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                para._element.insert(0, pPr)

            # 移除背景色
            shd = pPr.find(qn("w:shd"))
            if shd is not None:
                pPr.remove(shd)

            # 移除边框
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)

            # 缩进
            para.paragraph_format.left_indent = Cm(1.5)
            # 上下间距
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

            # 在文本开头插入破折号
            text = para.text.strip()
            if text and not text.startswith("——"):
                first_run = para.runs[0] if para.runs else None
                if first_run:
                    first_run.text = "—— " + first_run.text

            # 所有 run 设为斜体 + 深灰色
            for run in para.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.font.size = Pt(9.5)
                run.font.name = FONT_CN
                _set_east_asia_font(run)

    print(f"  [3/5] 引用块调整完成（{count} 个）")


def step4_remove_horizontal_rules(doc: Document):
    """删除分隔线（pandoc 用 VML v:rect o:hr=t 或 pBdr 实现）"""
    count = 0
    paragraphs_to_remove = []

    VML_NS = "urn:schemas-microsoft-com:vml"
    OFFICE_NS = "urn:schemas-microsoft-com:office:office"

    for para in doc.paragraphs:
        el = para._element
        # 方式1: VML 图形 hr（v:rect with o:hr="t"）
        for pict in el.findall(f".//{{{VML_NS}}}rect"):
            if pict.get(f"{{{OFFICE_NS}}}hr") == "t":
                paragraphs_to_remove.append(para)
                count += 1
                break
        else:
            # 方式2: 带底边框的空段落
            pPr = el.find(qn("w:pPr"))
            if pPr is not None:
                pBdr = pPr.find(qn("w:pBdr"))
                if pBdr is not None:
                    bottom = pBdr.find(qn("w:bottom"))
                    if bottom is not None and not para.text.strip():
                        paragraphs_to_remove.append(para)
                        count += 1

    for para in paragraphs_to_remove:
        parent = para._element.getparent()
        parent.remove(para._element)

    print(f"  [4/5] 分隔线删除完成（{count} 条）")


def step6_normalize_images(doc: Document):
    """图片排版规范化（彻底重建版）：
    - 提取所有 inline 图片的 rId 和原始尺寸
    - 删除 pandoc 生成的 'Captioned Figure' 段落（容器有 floating 残留属性）
    - 用 python-docx add_picture 重新以 inline 方式插入到原位置
    - 等比缩放到版心宽度 14cm
    - "图 N：xxx" 图注段落格式化为居中、灰色、9pt
    """
    from docx.shared import Emu, Inches
    import re

    TARGET_WIDTH_CM = 14
    image_count = 0
    caption_count = 0

    WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    PIC_NS = "{http://schemas.openxmlformats.org/drawingml/2006/picture}"

    # 收集所有图片信息：(段落对象, rId, 原始 cx, 原始 cy)
    image_paragraphs = []
    for para in doc.paragraphs:
        for blip in para._element.iter(f"{A_NS}blip"):
            rId = blip.get(f"{R_NS}embed")
            extent = next(para._element.iter(f"{WP_NS}extent"), None)
            if rId and extent is not None:
                cx = int(extent.get("cx", 0))
                cy = int(extent.get("cy", 0))
                image_paragraphs.append((para, rId, cx, cy))
                break

    # 处理每张图：清空段落内容，用 python-docx 添加干净的 inline 图
    for para, rId, orig_cx, orig_cy in image_paragraphs:
        # 通过 rId 拿到图片实际文件名
        try:
            image_part = doc.part.related_parts[rId]
            image_blob = image_part.blob
        except KeyError:
            continue

        # 等比缩放
        if orig_cx > 0:
            target_cx_emu = Cm(TARGET_WIDTH_CM).emu
            if orig_cx > target_cx_emu:
                width = Cm(TARGET_WIDTH_CM)
            else:
                width = Emu(orig_cx)
        else:
            width = Cm(TARGET_WIDTH_CM)

        # 清空原段落所有 run
        for r in list(para._element.findall(qn("w:r"))):
            para._element.remove(r)

        # 段落级属性：居中、无首行缩进、上下间距
        # 行距：单倍 + 清除固定行距设定（让段落高度自适应图片高度，避免图被裁切）
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        # 显式清除"固定行距"规则，让 Word 按 auto 处理
        from docx.enum.text import WD_LINE_SPACING
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 用 python-docx 加 inline 图
        import io
        run = para.add_run()
        run.add_picture(io.BytesIO(image_blob), width=width)

        image_count += 1

    # 处理图注段落：全文搜索以"图"开头且紧跟编号或冒号的段落
    # 匹配格式：
    #   "图：xxx"、"图:xxx"（无编号）
    #   "图 N：xxx"、"图 N. xxx"（带数字编号）
    #   "图 一：xxx"（带中文数字编号）
    from docx.enum.text import WD_LINE_SPACING
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r"^图\s*([\d一二三四五六七八九十]+)?\s*[：:\.\s]", text):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.2
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = False
                run.font.bold = False
                run.font.name = FONT_CN
                _set_east_asia_font(run)
            caption_count += 1

    print(f"  [6/6] 图片规范化完成（{image_count} 张图，{caption_count} 个图注）")


def step5_normalize_layout(doc: Document):
    """全局排版规范化：字体统一、正文字号/行距、首行缩进、列表缩进"""
    count = 0
    indent_count = 0
    # 不需要首行缩进的样式前缀
    NO_INDENT_STYLES = ("Heading", "Title", "Subtitle", "TOC", "List", "Block Text",
                        "Caption", "Header", "Footer", "Compact")
    # 正文样式（需要设定字号和行距）
    BODY_STYLES = ("Normal", "Body Text", "First Paragraph", "Body")

    for para in doc.paragraphs:
        style_name = para.style.name or ""

        # 所有 run 统一字体
        for run in para.runs:
            run.font.name = FONT_CN
            _set_east_asia_font(run)
            count += 1

        # 正文段落：字号12pt（小四）、行距28磅、首行缩进
        if any(style_name.startswith(s) for s in BODY_STYLES):
            for run in para.runs:
                if run.font.size is None or run.font.size == Pt(0):
                    run.font.size = Pt(12)
            para.paragraph_format.line_spacing = Pt(28)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if para.text.strip():
                para.paragraph_format.first_line_indent = Pt(24)  # 2×12pt
                indent_count += 1

        # 列表项：压缩左缩进 + 行距28磅
        elif style_name.startswith("List") or style_name.startswith("Compact"):
            para.paragraph_format.left_indent = Cm(0.8)
            para.paragraph_format.line_spacing = Pt(28)

        # 标题：确保行距28磅
        elif style_name.startswith("Heading"):
            para.paragraph_format.line_spacing = Pt(28)

    # 表格单元格字体统一（step2 已处理样式，这里确保无遗漏）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_CN
                        _set_east_asia_font(run)
                        count += 1
    print(f"  [5/5] 排版规范化完成（{count} 个文本块，{indent_count} 段首行缩进）")


# --- 辅助函数 ---

# 版心宽度 = 21cm - 3.17cm × 2 = 14.66cm ≈ 8314 twips (1cm = 567 twips)
PAGE_CONTENT_WIDTH_TWIPS = 8314


def _auto_column_widths(table):
    """根据每列最大内容长度智能分配列宽（短列紧凑，长列按比例分配）"""
    if not table.rows:
        return
    num_cols = len(table.columns)
    if num_cols == 0:
        return

    # 计算每列：表头长度（必须单行显示）和数据区最大长度
    header_lens = [0] * num_cols
    data_max_lens = [0] * num_cols
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= num_cols:
                break
            text = cell.text.strip()
            char_len = sum(2 if ord(c) > 0x7F else 1 for c in text)
            if row_idx == 0:
                header_lens[col_idx] = char_len
            else:
                data_max_lens[col_idx] = max(data_max_lens[col_idx], char_len)

    # 每列的有效宽度 = 至少能放下表头（不换行），同时参考数据区宽度
    max_lens = [max(h, d) for h, d in zip(header_lens, data_max_lens)]

    # 9pt 微软雅黑加粗：中文字符约 200twips宽
    CHAR_TWIPS = 200      # 每字符宽度（含间距余量）
    PADDING_TWIPS = 300   # 单元格左右内边距 + 余量
    MIN_COL_TWIPS = 800   # 最小列宽约 1.4cm
    MAX_SHORT_LEN = 12    # ≤12 字符宽度的列视为"短列"

    # 短列：保证表头不换行；长列：按比例分配
    col_widths = [0] * num_cols
    short_total = 0
    long_cols = []

    for i in range(num_cols):
        # 短列宽度至少要放下表头
        header_w = header_lens[i] * CHAR_TWIPS + PADDING_TWIPS
        if max_lens[i] <= MAX_SHORT_LEN:
            w = max(max_lens[i] * CHAR_TWIPS + PADDING_TWIPS, header_w, MIN_COL_TWIPS)
            col_widths[i] = w
            short_total += w
        else:
            long_cols.append(i)

    # 长列按比例分配剩余空间
    remaining = PAGE_CONTENT_WIDTH_TWIPS - short_total
    if remaining < MIN_COL_TWIPS * len(long_cols):
        # 空间不够，全部平分
        avg = PAGE_CONTENT_WIDTH_TWIPS // num_cols
        col_widths = [avg] * num_cols
    elif long_cols:
        long_total_len = sum(max_lens[i] for i in long_cols)
        for i in long_cols:
            ratio = max_lens[i] / long_total_len if long_total_len > 0 else 1 / len(long_cols)
            col_widths[i] = max(int(remaining * ratio), MIN_COL_TWIPS)

    # 应用列宽到每行的每个单元格
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= num_cols:
                break
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                cell._tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:type="dxa" w:w="{col_widths[col_idx]}"/>')
                tcPr.append(tcW)
            else:
                tcW.set(qn("w:type"), "dxa")
                tcW.set(qn("w:w"), str(col_widths[col_idx]))


def _set_cell_shading(tcPr, fill_color: str):
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_color}"/>'
    ))


def _set_east_asia_font(run):
    """设置 run 的东亚字体为中文字体"""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_CN}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), FONT_CN)


def _format_cell_runs(cell, color: RGBColor, bold: bool):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = color
            run.font.bold = bold
            run.font.size = Pt(9)
            run.font.name = FONT_CN
            rPr = run._element.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = parse_xml(
                        f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_CN}"/>'
                    )
                    rPr.insert(0, rFonts)
                else:
                    rFonts.set(qn("w:eastAsia"), FONT_CN)


def main():
    if len(sys.argv) < 2:
        print("用法: python md2docx.py input.md [output.docx]")
        sys.exit(1)

    md_path = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        docx_path = str(Path(md_path).with_suffix(".docx"))

    print(f"转换: {md_path}")
    print(f"输出: {docx_path}")

    # Step 1: pandoc 转换
    step1_pandoc_convert(md_path, docx_path)

    # Step 2-5: python-docx 后处理
    doc = Document(docx_path)
    step2_beautify_tables(doc)
    step3_fix_blockquotes(doc)
    step4_remove_horizontal_rules(doc)
    step5_normalize_layout(doc)
    step6_normalize_images(doc)
    doc.save(docx_path)

    print(f"完成: {docx_path}")


if __name__ == "__main__":
    main()
